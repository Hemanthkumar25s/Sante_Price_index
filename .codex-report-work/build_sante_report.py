from copy import deepcopy
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
WORK = ROOT / ".codex-report-work"
INPUT = WORK / "intership report.docx"
OUTPUT = WORK / "Sante_Price_Index_Internship_Report_Clean.docx"
SCREENSHOTS = ROOT / "screenshots"
ASSETS = WORK / "assets"

ACCENT = RGBColor(31, 78, 121)
MUTED = RGBColor(89, 89, 89)
LIGHT_FILL = "EAF2F8"


def set_cell_text(cell, text, bold=False, size=12, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table, color="B7C9D9"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_cell_margins(table, top=120, start=120, bottom=120, end=120):
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn("w:" + m))
        if node is None:
            node = OxmlElement("w:" + m)
            margins.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_horizontal_line(paragraph, color="4F81BD", size="12"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "8")
    bottom.set(qn("w:color"), color)


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    styles["Normal"].font.size = Pt(12)

    for name, size in (("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 14)):
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = ACCENT
        st.paragraph_format.space_before = Pt(12)
        st.paragraph_format.space_after = Pt(6)

    for p in doc.paragraphs:
        style_name = p.style.name if p.style else ""
        if style_name == "Heading 1":
            target_size = 16
            bold = True
            color = ACCENT
        elif style_name in ("Heading 2", "Heading 3"):
            target_size = 14
            bold = True
            color = ACCENT
        else:
            target_size = 12
            bold = None
            color = None
        for run in p.runs:
            set_run_font(run, size=target_size, bold=bold, color=color)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(6)

    for table in doc.tables:
        set_table_borders(table)
        set_cell_margins(table)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row_i, row in enumerate(table.rows):
            for cell in row.cells:
                if row_i == 0:
                    shade_cell(cell, LIGHT_FILL)
                for p in cell.paragraphs:
                    p.paragraph_format.line_spacing = 1.05
                    p.paragraph_format.space_after = Pt(3)
                    for run in p.runs:
                        set_run_font(run, size=12, bold=(row_i == 0), color=ACCENT if row_i == 0 else None)


def remove_existing_chapter_6(doc):
    body = doc.element.body
    children = list(body)
    start = None
    for idx, child in enumerate(children):
        if child.tag == qn("w:p"):
            text = "".join(node.text or "" for node in child.iter() if node.tag == qn("w:t")).strip()
            if (text.startswith("6.") and "Technical Attachments" in text) or (
                text.startswith("6.") and "Project Details" in text
            ):
                start = idx
                break
    if start is None:
        return
    for child in children[start:]:
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size=16 if level == 1 else 14, bold=True, color=ACCENT)
    return p


def add_body_para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, size=12)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.15)
        p.paragraph_format.line_spacing = 1.1
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run("- " + item)
        set_run_font(run, size=12)


def jpeg_asset(source):
    ASSETS.mkdir(exist_ok=True)
    target = ASSETS / (source.stem + ".jpg")
    if not target.exists():
        image = Image.open(source).convert("RGB")
        image.save(target, "JPEG", quality=92, optimize=True)
    return target


def add_simple_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    set_cell_margins(table)
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, color=ACCENT)
        shade_cell(table.rows[0].cells[i], LIGHT_FILL)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = width
    doc.add_paragraph()
    return table


def add_chapter_6(doc):
    add_heading(doc, "6. Project Details and Screenshots", 1)
    add_body_para(
        doc,
        "The Sante Price Index is an Android application developed for fresh-market vendors, retailers, "
        "and small produce sellers who need quick access to mandi prices, pricing calculations, stock "
        "visibility, and customer-facing price-board information. The project combines real-time data, "
        "Jetpack Compose screens, Firebase integration, and a local AI assistant to support practical "
        "daily trading decisions."
    )

    add_heading(doc, "Project Overview", 2)
    add_simple_table(
        doc,
        ["Item", "Project Information"],
        [
            ["Application Name", "Sante Price Index"],
            ["Package Name", "com.sante.priceindex"],
            ["Platform", "Android mobile application"],
            ["Minimum SDK", "Android 8.0, API 26"],
            ["Target SDK", "Android API 35"],
            ["Architecture", "MVVM with Repository pattern and reactive StateFlow state management"],
            ["Primary Users", "Vegetable vendors, small retailers, staff users, and shop owners"],
        ],
        [Inches(1.8), Inches(4.5)],
    )

    add_heading(doc, "Technical Stack", 2)
    add_simple_table(
        doc,
        ["Component", "Technology Used"],
        [
            ["Programming Language", "Kotlin"],
            ["User Interface", "Jetpack Compose with Material 3"],
            ["Navigation", "Jetpack Compose Navigation"],
            ["Backend", "Firebase Realtime Database"],
            ["Authentication", "Firebase Authentication and Google Sign-In with Credential Manager"],
            ["State Management", "Kotlin Coroutines, Flow, and StateFlow"],
            ["Local Storage", "DataStore Preferences"],
            ["Image Loading", "Coil Compose"],
            ["Build System", "Gradle Kotlin DSL"],
            ["Development Tool", "Android Studio"],
        ],
        [Inches(2.0), Inches(4.3)],
    )

    add_heading(doc, "Core Modules Implemented", 2)
    add_simple_table(
        doc,
        ["Module", "Description"],
        [
            ["Login and Authentication", "Email/password and Google Sign-In based onboarding with protected app navigation."],
            ["Home Dashboard", "Central navigation area that summarizes market movement, quick actions, and vendor workflow entry points."],
            ["Price Watch", "Live mandi-price listing with price movement indicators and commodity selection."],
            ["Profit Calculator", "Computes cost per kilogram, retail price, gross sales, and net profit using quantity, transport, wastage, and margin inputs."],
            ["Digital Price Board", "Customer-facing price display that can be updated from calculated prices."],
            ["Market Trends", "Seven-day price history visualization to support short-term price awareness."],
            ["Smart Alerts", "Highlights rising and falling commodity prices for faster vendor decisions."],
            ["AI Assistant", "Local query assistant for produce-price questions, including transliterated Kannada and Hindi examples."],
            ["Profile and Roles", "Vendor profile fields, language selection, and role-based access for admin, vendor, and staff users."],
        ],
        [Inches(1.8), Inches(4.6)],
    )

    add_heading(doc, "Implementation Highlights", 2)
    add_bullets(
        doc,
        [
            "The application starts with a splash and authentication flow, then routes logged-in users to the main dashboard.",
            "A bottom navigation bar gives quick access to Home, Prices, AI Agent, Calculator, and Board screens.",
            "MainViewModel acts as the central state holder for prices, selected commodities, profit inputs, board items, inventory, role, and language.",
            "FirebaseRepository listens to the mandi_prices database node and falls back to seed data when live data is temporarily unavailable.",
            "The profit calculator uses a transparent formula based on mandi cost, transport cost, wastage buffer, margin percentage, and quantity.",
            "Role-based restrictions keep staff users focused on the Price Board while vendor/admin users can access calculation and AI features.",
        ],
    )

    add_heading(doc, "Screenshot Gallery", 2)
    add_body_para(
        doc,
        "The following screenshots from the project folder show the important user flows and completed screens of the Sante Price Index application."
    )

    shots = [
        ("login.png", "Login screen", "Secure sign-in entry point for vendor access."),
        ("home.png", "Home dashboard", "Quick navigation to market tools and daily price activity."),
        ("price_watch.png", "Price Watch", "Live commodity-price list with market update context."),
        ("profit_calculator.png", "Profit Calculator", "Input-based calculation of cost, margin, and retail selling price."),
        ("board.png", "Digital Price Board", "Customer-facing board for the shop's current selling prices."),
        ("price_trends.png", "Price Trends", "Seven-day movement view for commodity price analysis."),
        ("best_market.png", "Best Market", "Decision support for identifying better selling opportunities."),
        ("smart_alerts.png", "Smart Alerts", "Alerts for important price changes and market movement."),
        ("Ai_chat_bot.png", "Sante AI Agent", "Chat interface for produce-business and price questions."),
        ("profile.png", "Profile screen", "Vendor profile, shop details, and preference management."),
    ]

    for idx, shot in enumerate(shots):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        filename, title, caption = shot
        image_path = SCREENSHOTS / filename
        if image_path.exists():
            p.add_run().add_picture(str(jpeg_asset(image_path)), width=Inches(2.65))

        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(f"Figure {idx + 1}: {title}. ")
        set_run_font(run, size=12, bold=True, color=ACCENT)
        run2 = cap.add_run(caption)
        set_run_font(run2, size=12, color=MUTED)
        if idx != len(shots) - 1:
            cap.add_run().add_break(WD_BREAK.PAGE)

    add_heading(doc, "Testing and Validation", 2)
    add_body_para(
        doc,
        "The project includes a unit-test scaffold and was validated through Android Studio builds, emulator/device execution, "
        "navigation checks, Firebase data fallback checks, profit-calculation verification, and manual review of the major UI screens. "
        "The current implementation keeps the business logic readable so that future automated tests can be added for repositories, "
        "view models, and calculation rules."
    )

    add_heading(doc, "Future Enhancements", 2)
    add_bullets(
        doc,
        [
            "Replace seed demo data with a fully managed production Firebase database and admin data-entry workflow.",
            "Add push notifications for sharp price changes and stock-level warnings.",
            "Introduce offline caching so vendors can continue viewing recent prices when network access is weak.",
            "Expand the AI assistant with richer multilingual support and voice-first interactions for local vendors.",
            "Add exportable daily sales and stock reports for shop accounting.",
        ],
    )

    add_heading(doc, "Chapter Summary", 2)
    add_body_para(
        doc,
        "This chapter documents the actual Sante Price Index project implementation, including its architecture, tools, "
        "feature modules, UI evidence, and future scope. The screenshots confirm that the internship work moved beyond "
        "conceptual design into a complete Android application with multiple practical vendor-facing workflows."
    )


def add_front_matter(doc):
    body = doc.element.body
    existing_elements = set(list(body))

    def mark_new(func, *args, **kwargs):
        return func(*args, **kwargs)

    p = mark_new(doc.add_paragraph)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("INTERNSHIP PROJECT REPORT")
    set_run_font(r, size=16, bold=True, color=ACCENT)

    for _ in range(2):
        mark_new(doc.add_paragraph)

    p = mark_new(doc.add_paragraph)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Sante Price Index")
    set_run_font(r, size=16, bold=True, color=ACCENT)

    p = mark_new(doc.add_paragraph)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("An Android Application for Fresh-Market Price Intelligence")
    set_run_font(r, size=14, bold=True, color=MUTED)

    for text in [
        "Submitted in partial fulfilment of the requirements for the",
        "Bachelor of Engineering - Information Science and Engineering",
        "Internship Organisation",
        "MindMatrix.io",
        "Kiran Arcade, 4th Floor, Sector 1, HSR Layout, Bengaluru, Karnataka",
        "Programme: Android App Development Using Generative AI",
        "JSS Academy of Technical Education, Bengaluru (JSSATEB)",
        "Department of Information Science and Engineering",
        "Academic Year 2024-25",
    ]:
        p = mark_new(doc.add_paragraph)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        if text in {"Internship Organisation", "Programme: Android App Development Using Generative AI"}:
            set_run_font(r, size=12, bold=True, color=ACCENT)
        else:
            set_run_font(r, size=12)

    p = mark_new(doc.add_paragraph)
    p.add_run().add_break(WD_BREAK.PAGE)

    p = mark_new(doc.add_paragraph)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("INDEX")
    set_run_font(r, size=16, bold=True, color=ACCENT)
    add_horizontal_line(p)

    p = mark_new(doc.add_paragraph)
    add_horizontal_line(p, color="4F81BD", size="8")

    toc_rows = [
        ("1.", "About the Organisation", "1"),
        ("1.1", "Background and Overview of MindMatrix.io", "1"),
        ("1.2", "Vision and Mission", "2"),
        ("1.3", "Core Values", "2"),
        ("1.4", "Organisational Milestones", "3"),
        ("1.5", "The 7P Framework", "4"),
        ("2.", "Objective of Internship", "5"),
        ("2.1", "Primary Objectives", "5"),
        ("2.2", "Secondary Objectives", "6"),
        ("3.", "Learning Experiences", "7"),
        ("3.1", "Jetpack Compose and Declarative UI Development", "7"),
        ("3.2", "Firebase Realtime Database and Authentication", "8"),
        ("3.3", "MVVM Architecture and Advanced State Management", "9"),
        ("3.4", "AI Repository and Multi-lingual NLP Logic", "10"),
        ("3.5", "Performance Optimisation and Startup Profiling", "11"),
        ("3.6", "Profit Optimiser and Digital Price Board", "12"),
        ("4.", "Learning Outcomes", "14"),
        ("4.1", "Technical Outcomes", "14"),
        ("4.2", "Professional Outcomes", "15"),
        ("5.", "Conclusion / Summary", "16"),
        ("5.1", "Summary of Work Done", "16"),
        ("5.2", "Key Takeaways", "17"),
        ("5.3", "Future Scope", "18"),
        ("5.4", "Final Reflection", "19"),
        ("6.", "Project Details and Screenshots", "20"),
        ("6.1", "Project Overview", "20"),
        ("6.2", "Technical Stack", "21"),
        ("6.3", "Core Modules Implemented", "22"),
        ("6.4", "Screenshot Gallery", "23"),
        ("6.5", "Testing, Validation, and Future Enhancements", "28"),
    ]
    table = mark_new(doc.add_table, rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    set_cell_margins(table, top=90, start=120, bottom=90, end=120)
    headers = ["Sl. No.", "Contents", "Page No."]
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, color=ACCENT)
        shade_cell(table.rows[0].cells[i], LIGHT_FILL)
    for number, title, page in toc_rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], number)
        set_cell_text(cells[1], title)
        set_cell_text(cells[2], page)
        cells[0].width = Inches(0.8)
        cells[1].width = Inches(4.6)
        cells[2].width = Inches(0.9)

    p = mark_new(doc.add_paragraph)
    p.add_run().add_break(WD_BREAK.PAGE)

    created = [el for el in body if el not in existing_elements and el.tag != qn("w:sectPr")]
    sect = body.sectPr
    for el in created:
        if el.getparent() is body:
            body.remove(el)
    insert_at = 0
    for el in created:
        body.insert(insert_at, el)
        insert_at += 1
    if sect is not None and sect.getparent() is None:
        body.append(sect)


def main():
    doc = Document(INPUT)
    style_document(doc)
    remove_existing_chapter_6(doc)
    add_chapter_6(doc)
    add_front_matter(doc)
    style_document(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
